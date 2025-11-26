# sudo.py
# -------------------------------------------------------------
# Smart Packing Assistant
# DO NOT store or hardcode API keys here.
# Use the .env file in the project root for environment variables.
# -------------------------------------------------------------

import os
from dotenv import load_dotenv
import streamlit as st
import requests
import google.generativeai as genai
from groq import Groq
from datetime import datetime
from io import BytesIO
import re
from docx import Document

# --- Load environment variables ---
load_dotenv()  # Reads .env in project root

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OPENCAGE_API_KEY = os.getenv("OPENCAGE_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)

# --- Streamlit Page Config ---
st.set_page_config(page_title="Smart Packing Assistant", page_icon="🎒", layout="wide")

# --- Cache for coordinates ---
location_cache = {}

def get_lat_lon(location):
    """Fetch latitude and longitude from OpenCage with fallback."""
    if location in location_cache:
        return location_cache[location]
    try:
        url = f"https://api.opencagedata.com/geocode/v1/json?q={location}&key={OPENCAGE_API_KEY}"
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        data = response.json()
        if data["results"]:
            lat = data["results"][0]["geometry"]["lat"]
            lon = data["results"][0]["geometry"]["lng"]
            location_cache[location] = (lat, lon)
            return lat, lon
    except Exception as e:
        st.warning(f"Could not fetch coordinates: {e}. Using fallback coordinates for testing.")
    # Fallback: New Delhi
    return 28.6139, 77.2090


def create_docx(weather, packing_list):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Smart Packing Assistant", level=1)
    doc.add_paragraph(f"Weather Forecast: {weather}\n")
    doc.add_heading("Packing List:", level=2)
    for item in packing_list:
        para = doc.add_paragraph("- ")
        parts = re.finditer(r"\*([^*]+)\*", item)
        last_index = 0
        for match in parts:
            para.add_run(item[last_index:match.start()])
            run = para.add_run(match.group(1))
            run.bold = True
            last_index = match.end()
        para.add_run(item[last_index:])
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def extract_items(text):
    return [re.sub(r"[:•\-]+", "", line).strip() for line in text.split("\n") if line.strip()]

def get_llama_suggestions(location, activities, people):
    """
    Get a structured packing list from LLaMA (Groq API) with bullet points only.
    """
    try:
        client = Groq(api_key=GROQ_API_KEY)
        query = (
            f"Suggest a personalized packing list for {len(people)} people traveling to {location} for {activities}. "
            f"Include age, gender, medical needs."
        )
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": query}],
            temperature=1,
            max_completion_tokens=300,
            top_p=1,
            stream=False
        )

        # Fixed content extraction for Groq SDK
        choice = completion.choices[0]
        if hasattr(choice, "message") and hasattr(choice.message, "content"):
            text = choice.message.content
        elif hasattr(choice, "message_text"):
            text = choice.message_text
        else:
            text = ""

        return extract_items(text)

    except Exception as e:
        st.warning(f"LLaMA API failed: {e}")
        return []

# --- Streamlit UI ---
st.title("🎒 Smart Packing Assistant")
st.markdown("### Generate personalized packing lists based on location, activities, and trip type!")

# Sidebar
st.sidebar.header("🌍 Trip Details")
location = st.sidebar.text_input("Enter a location:")
start_date = st.sidebar.date_input("Start Date", datetime.today())
end_date = st.sidebar.date_input("End Date", datetime.today())
trip_type = st.sidebar.selectbox("Select Trip Type", ["Business", "Leisure", "Adventure", "Other"])
activities = st.sidebar.text_area("Enter your activities (comma-separated):")
packing_list_type = st.sidebar.selectbox("Select Packing List Type", ["Detailed", "Minimal"])
num_people = st.sidebar.number_input("Number of travelers:", min_value=1, step=1)

# Traveler details
people = []
for i in range(num_people):
    with st.expander(f"👤 Person {i + 1} Details"):
        name = st.text_input(f"Name", key=f"name_{i}")
        gender = st.selectbox("Gender", ["Male", "Female", "Other"], key=f"gender_{i}")
        age = st.number_input("Age", min_value=0, step=1, key=f"age_{i}")
        medical = st.text_area("Medical Issues", key=f"medical_{i}")
        people.append({"name": name, "gender": gender, "age": age, "medical_issues": medical})

# Generate packing list
if st.sidebar.button("Generate Packing List 🧳"):
    if not location or not activities or not people:
        st.error("🚨 Please fill all required fields.")
    else:
        lat, lon = get_lat_lon(location)
        weather = "Sunny 25°C"  # Replace with actual weather API if needed

        # Use LLaMA directly, no Gemini
        packing_list = get_llama_suggestions(location, activities, people)

        st.success("✅ Packing List Generated!")
        st.subheader("🌤 Weather Forecast")
        st.info(weather)
        st.subheader(f"📦 {packing_list_type} Packing List")
        for item in packing_list:
            st.write(f"- {item}")

        docx_file = create_docx(weather, packing_list)
        st.download_button(
            "📥 Download DOCX Packing List",
            data=docx_file,
            file_name="Packing_List.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
